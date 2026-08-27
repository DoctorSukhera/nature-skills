from __future__ import annotations
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import json
import zipfile
import re

from docx import Document
from pypdf import PdfReader
import pandas as pd


@dataclass
class ManuscriptBundle:
    filename: str
    text: str
    tables: list[dict] = field(default_factory=list)
    images: list[dict] = field(default_factory=list)
    source_bytes: bytes | None = None


def _docx_text_and_tables(data: bytes) -> tuple[str, list[dict]]:
    doc = Document(BytesIO(data))
    chunks: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    tables: list[dict] = []
    for i, table in enumerate(doc.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append({'index': i, 'rows': rows})
            chunks.append(f'\n[TABLE {i}]\n' + '\n'.join(' | '.join(r) for r in rows))
    return '\n'.join(chunks), tables


def _docx_images(data: bytes) -> list[dict]:
    import xml.etree.ElementTree as ET
    images: list[dict] = []
    with zipfile.ZipFile(BytesIO(data)) as zf:
        try:
            rels_root = ET.fromstring(zf.read('word/_rels/document.xml.rels'))
            rels = {}
            for rel in rels_root:
                rid = rel.attrib.get('Id')
                target = rel.attrib.get('Target', '')
                if rid and 'media/' in target:
                    rels[rid] = 'word/' + target.lstrip('/')
            doc_root = ET.fromstring(zf.read('word/document.xml'))
            ns = {'a':'http://schemas.openxmlformats.org/drawingml/2006/main','r':'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            seen = set()
            for blip in doc_root.findall('.//a:blip', ns):
                rid = blip.attrib.get('{%s}embed' % ns['r'])
                name = rels.get(rid)
                if not name or name in seen or name not in zf.namelist():
                    continue
                seen.add(name)
                blob = zf.read(name)
                ext = Path(name).suffix.lower().lstrip('.') or 'png'
                images.append({'name': Path(name).name, 'ext': ext, 'bytes': blob})
            if images:
                return images
        except Exception:
            pass
        for name in sorted(zf.namelist()):
            if name.startswith('word/media/'):
                blob = zf.read(name)
                ext = Path(name).suffix.lower().lstrip('.') or 'png'
                images.append({'name': Path(name).name, 'ext': ext, 'bytes': blob})
    return images


def parse_upload(uploaded) -> ManuscriptBundle:
    data = uploaded.getvalue()
    name = uploaded.name
    suffix = Path(name).suffix.lower()
    if suffix == '.docx':
        text, tables = _docx_text_and_tables(data)
        return ManuscriptBundle(name, text, tables, _docx_images(data), data)
    if suffix == '.pdf':
        reader = PdfReader(BytesIO(data))
        text = '\n'.join((page.extract_text() or '') for page in reader.pages)
        return ManuscriptBundle(name, text, [], [], data)
    if suffix in {'.txt', '.md'}:
        return ManuscriptBundle(name, data.decode('utf-8', errors='replace'), [], [], data)
    raise ValueError('Supported manuscript formats: DOCX, PDF, TXT, Markdown.')


def _support_text(uploaded, per_file_limit: int = 26000) -> str:
    data = uploaded.getvalue()
    name = uploaded.name
    suffix = Path(name).suffix.lower()
    head = f'\n\n===== SUPPORTING FILE: {name} =====\n'
    try:
        if suffix == '.docx':
            text, tables = _docx_text_and_tables(data)
            return head + text[:per_file_limit]
        if suffix == '.pdf':
            reader = PdfReader(BytesIO(data))
            text = '\n'.join((p.extract_text() or '') for p in reader.pages)
            return head + text[:per_file_limit]
        if suffix in {'.txt','.md','.py','.r','.json','.yaml','.yml','.csv'}:
            if suffix == '.csv':
                df = pd.read_csv(BytesIO(data), nrows=200)
                text = df.to_csv(index=False)
            else:
                text = data.decode('utf-8', errors='replace')
                if suffix == '.json':
                    try: text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
                    except Exception: pass
            return head + text[:per_file_limit]
        if suffix in {'.xlsx','.xls'}:
            xl = pd.ExcelFile(BytesIO(data))
            parts=[]
            for sheet in xl.sheet_names[:8]:
                df = pd.read_excel(xl, sheet_name=sheet, nrows=200)
                parts.append(f'[SHEET {sheet}]\n'+df.to_csv(index=False))
            return head + '\n'.join(parts)[:per_file_limit]
        if suffix == '.ipynb':
            nb = json.loads(data.decode('utf-8', errors='replace'))
            parts=[]
            for cell in nb.get('cells',[]):
                src=''.join(cell.get('source',[]))
                if src.strip(): parts.append(f"[{cell.get('cell_type','cell')}]\n{src}")
            return head + '\n\n'.join(parts)[:per_file_limit]
    except Exception as exc:
        return head + f'[Could not parse supporting file: {exc}]'
    return head + '[Unsupported supporting-file format]'


def parse_support_uploads(files, total_limit: int = 110000) -> str:
    if not files:
        return ''
    blocks=[]; used=0
    for f in files:
        block = _support_text(f)
        if used + len(block) > total_limit:
            block = block[:max(0,total_limit-used)]
        if block:
            blocks.append(block); used += len(block)
        if used >= total_limit:
            break
    return ''.join(blocks)


def count_citations(text: str) -> int:
    return len(re.findall(r'\[(?:\d+(?:\s*[-–,]\s*\d+)*)\]', text))
