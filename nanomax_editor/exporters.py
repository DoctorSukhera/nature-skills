from __future__ import annotations
from io import BytesIO
import json
import re
import zipfile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from PIL import Image, ImageOps, ImageDraw
except Exception:
    Image = None

CITE_RE = re.compile(r"\[\[CITE:([0-9,\-–\s]+)\]\]")


def _clean_inline_text(text: str) -> str:
    """Remove model-facing markup that must never appear in a manuscript."""
    text = str(text or "")
    text = text.replace("\\<", "<").replace("\\>", ">")
    text = re.sub(r"</?(?:b|strong|i|em|span)[^>]*>", "", text, flags=re.I)
    text = re.sub(r"\\</?(?:b|strong|i|em|span)[^>]*\\>", "", text, flags=re.I)
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"^#{1,6}\s+", "", text.strip())
    return text.strip()


def _clean_pkg_in_place(pkg: dict) -> dict:
    """Deterministic last-line manuscript hygiene. Does not invent or rewrite science."""
    banned = ("nanomax nature editor", "not affiliated with or endorsed by nature portfolio")
    for sec_key in ("main_sections", "methods_sections", "supplementary_sections"):
        for sec in pkg.get(sec_key, []) or []:
            sec["heading"] = _clean_inline_text(sec.get("heading", ""))
            cleaned=[]
            for para in sec.get("paragraphs", []) or []:
                para = _clean_inline_text(para)
                if para and not any(x in para.lower() for x in banned):
                    cleaned.append(para)
            sec["paragraphs"] = cleaned
    for key in ("title","short_title","abstract","methods_parent_heading","data_availability","code_availability","ethics_statement","author_contributions","competing_interests","funding_statement","acknowledgements"):
        if key in pkg:
            pkg[key] = _clean_inline_text(pkg.get(key, ""))
    fm=pkg.get("front_matter") or {}
    for key in ("authors_line","correspondence","equal_contribution"):
        if key in fm: fm[key]=_clean_inline_text(fm.get(key,""))
    fm["affiliations"]=[_clean_inline_text(x) for x in fm.get("affiliations",[]) or []]
    fm["other_notes"]=[_clean_inline_text(x) for x in fm.get("other_notes",[]) or [] if not any(b in str(x).lower() for b in banned)]
    for f in pkg.get("figure_plan",[]) or []:
        for key in ("figure_number","title","legend","placement_after","reason","redraw_prompt"):
            f[key]=_clean_inline_text(f.get(key,""))
        f["content_lock"]=[_clean_inline_text(x) for x in f.get("content_lock",[]) or []]
    for t in pkg.get("tables",[]) or []:
        for key in ("number","title","footnote","placement_after"):
            t[key]=_clean_inline_text(t.get(key,""))
        t["columns"]=[_clean_inline_text(x) for x in t.get("columns",[]) or []]
        t["rows"]=[[_clean_inline_text(x) for x in row] for row in t.get("rows",[]) or []]
    pkg["references"]=[_clean_inline_text(x) for x in pkg.get("references",[]) or []]
    return pkg


def _style_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size in [("Title",18),("Heading 1",13),("Heading 2",11.5),("Heading 3",10.5)]:
        try:
            st = doc.styles[style_name]
            st.font.name = "Arial"; st.font.size = Pt(size); st.font.bold = True
        except Exception:
            pass


def _shade_cell(cell, fill="EAF0F7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def _add_citation_rich_text(paragraph, text: str):
    text = _clean_inline_text(text)
    pos = 0
    for m in CITE_RE.finditer(text or ""):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        cite = m.group(1).replace("–", "-").strip()
        r = paragraph.add_run(cite)
        r.font.superscript = True
        pos = m.end()
    if pos < len(text or ""):
        paragraph.add_run((text or "")[pos:])


def _add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_citation_rich_text(p, text)
    return p


def _safe_png(blob: bytes) -> bytes:
    if Image is None:
        return blob
    try:
        im = Image.open(BytesIO(blob))
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGBA" if "A" in im.getbands() else "RGB")
        out = BytesIO(); im.save(out, format="PNG"); return out.getvalue()
    except Exception:
        return blob


def _compose_assets(blobs: list[bytes]) -> bytes | None:
    if not blobs:
        return None
    if len(blobs) == 1:
        return _safe_png(blobs[0])
    if Image is None:
        return _safe_png(blobs[0])
    ims = []
    for b in blobs:
        try:
            im = Image.open(BytesIO(b)).convert("RGB")
            ims.append(im)
        except Exception:
            continue
    if not ims:
        return None
    target_w = 1200
    thumbs = []
    for im in ims:
        ratio = target_w / max(im.width, 1)
        h = max(1, int(im.height * ratio))
        thumbs.append(im.resize((target_w, h)))
    pad = 24
    canvas_h = sum(im.height for im in thumbs) + pad*(len(thumbs)-1)
    canvas = Image.new("RGB", (target_w, canvas_h), "white")
    y = 0
    draw = ImageDraw.Draw(canvas)
    for i, im in enumerate(thumbs):
        canvas.paste(im, (0,y))
        draw.rectangle((4,y+4,42,y+42), fill="white")
        draw.text((13,y+9), chr(65+i), fill="black")
        y += im.height + pad
    out = BytesIO(); canvas.save(out, format="PNG"); return out.getvalue()


def _asset_blob(indices: list[int], source_images: list[dict], redrawn_images: dict[int, bytes] | None) -> bytes | None:
    blobs = []
    for idx in indices:
        if redrawn_images and idx in redrawn_images:
            blobs.append(redrawn_images[idx])
        elif 1 <= idx <= len(source_images):
            blobs.append(source_images[idx-1]["bytes"])
    return _compose_assets(blobs)


def _insert_table(doc: Document, t: dict):
    p = doc.add_paragraph()
    r = p.add_run(f"{t.get('number','')} | {t.get('title','')}")
    r.bold = True
    cols = t.get("columns") or []
    rows = t.get("rows") or []
    if cols:
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]; _set_repeat_table_header(hdr)
        for j,c in enumerate(cols):
            cell = hdr.cells[j]; cell.text = str(c); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; _shade_cell(cell)
            for rr in cell.paragraphs[0].runs: rr.bold = True
        for row in rows:
            cells = table.add_row().cells
            for j in range(len(cols)):
                cells[j].text = str(row[j]) if j < len(row) else ""
                cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if t.get("footnote"):
        p = doc.add_paragraph(); p.style = doc.styles["Normal"]
        r = p.add_run(t["footnote"]); r.italic = True; r.font.size = Pt(9)


def _insert_figure(doc: Document, f: dict, source_images: list[dict], redrawn_images: dict[int, bytes] | None):
    blob = _asset_blob(f.get("source_asset_indices") or [], source_images, redrawn_images)
    if blob:
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(); run.add_picture(BytesIO(blob), width=Inches(6.45))
        except Exception:
            doc.add_paragraph("[Figure image could not be embedded; use packaged source asset.]")
    else:
        doc.add_paragraph("[Figure source data/image required before submission.]")
    cap = doc.add_paragraph()
    number = f.get('figure_number','').strip()
    title = f.get('title','').strip()
    legend = f.get("legend", "").strip()
    # Prevent the model from duplicating a full figure title inside the legend.
    prefix_candidates = [f"{number} | {title}", f"{number}. {title}", f"{title}.", title]
    for pref in prefix_candidates:
        if pref and legend.lower().startswith(pref.lower()):
            legend = legend[len(pref):].lstrip(" .:|-–—")
            break
    r = cap.add_run(f"{number} | {title}. "); r.bold = True
    _add_citation_rich_text(cap, legend)
    cap.paragraph_format.space_after = Pt(9)


def _placement_matches(placement: str, heading: str) -> bool:
    a = (placement or "").strip().lower(); b = (heading or "").strip().lower()
    if not a or not b: return False
    return a == b or a in b or b in a


def master_manuscript_docx_bytes(pkg: dict, *, source_images: list[dict], redrawn_images: dict[int, bytes] | None = None, graphical_abstract: bytes | None = None, journal_profile: dict | None = None) -> bytes:
    """Build the primary all-in-one initial-submission Word manuscript with text, tables and figures embedded."""
    import copy
    pkg = _clean_pkg_in_place(copy.deepcopy(pkg))
    doc = Document(); _style_doc(doc)

    # Title/front matter
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(pkg.get("title", "")); r.bold = True; r.font.size = Pt(17)
    fm = pkg.get("front_matter") or {}
    if fm.get("authors_line"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(fm["authors_line"]).bold = True
    for aff in fm.get("affiliations", []) or []:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(aff).font.size = Pt(9.5)
    for txt in [fm.get("equal_contribution",""), fm.get("correspondence","")]:
        if txt:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(txt).font.size = Pt(9.5)
    for note in fm.get("other_notes",[]) or []:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(note).font.size = Pt(9)

    doc.add_heading("Abstract", level=1); _add_body_paragraph(doc, pkg.get("abstract", ""))
    kws = pkg.get("keywords") or []
    if kws:
        p = doc.add_paragraph(); p.add_run("Keywords: ").bold = True; p.add_run(", ".join(kws))

    ga = pkg.get("graphical_abstract") or {}
    if graphical_abstract and ga.get("include_in_manuscript"):
        doc.add_heading("Graphical abstract", level=1)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(_safe_png(graphical_abstract)), width=Inches(6.45))
        if ga.get("caption"): _add_body_paragraph(doc, ga["caption"])

    main_tables = [t for t in pkg.get("tables",[]) if t.get("destination") == "main"]
    main_figs = [f for f in pkg.get("figure_plan",[]) if f.get("destination") == "main" and f.get("action") != "remove_if_redundant"]
    used_t, used_f = set(), set()

    for si, sec in enumerate(pkg.get("main_sections", [])):
        heading = sec.get("heading", "")
        if sec.get("show_heading") and heading:
            doc.add_heading(heading, level=1)
        for para in sec.get("paragraphs", []) or []:
            _add_body_paragraph(doc, para)
        for i,t in enumerate(main_tables):
            if i not in used_t and _placement_matches(t.get("placement_after",""), heading):
                _insert_table(doc, t); used_t.add(i)
        for i,f in enumerate(main_figs):
            if i not in used_f and _placement_matches(f.get("placement_after",""), heading):
                _insert_figure(doc, f, source_images, redrawn_images); used_f.add(i)

    # Any unplaced main display items go after Results/main text, rather than being lost.
    for i,t in enumerate(main_tables):
        if i not in used_t: _insert_table(doc, t)
    for i,f in enumerate(main_figs):
        if i not in used_f: _insert_figure(doc, f, source_images, redrawn_images)

    if pkg.get("methods_sections"):
        parent_heading = pkg.get("methods_parent_heading", "Methods")
        if pkg.get("methods_parent_show_heading", True) and parent_heading:
            doc.add_heading(parent_heading, level=1)
        for sec in pkg.get("methods_sections", []):
            if sec.get("show_heading") and sec.get("heading"):
                level = max(1, min(3, int(sec.get("heading_level", 2) or 2)))
                # Methods topical subheadings are subordinate to the parent heading.
                if pkg.get("methods_parent_show_heading", True): level = max(2, level)
                doc.add_heading(sec["heading"], level=level)
            for para in sec.get("paragraphs",[]) or []: _add_body_paragraph(doc, para)

    role_map = {
        "data_availability": ("Data availability", "data_availability"),
        "code_availability": ("Code availability", "code_availability"),
        "ethics": ("Ethics and governance", "ethics_statement"),
        "funding": ("Funding", "funding_statement"),
        "acknowledgements": ("Acknowledgements", "acknowledgements"),
        "author_contributions": ("Author contributions", "author_contributions"),
        "competing_interests": ("Competing interests", "competing_interests"),
    }
    rendered_roles=set()
    outline = sorted((journal_profile or {}).get("outline_contract",[]) or [], key=lambda x:x.get("position",999))
    for item in outline:
        role=item.get("role")
        if role in role_map and role not in rendered_roles:
            default_label,key=role_map[role]; val=pkg.get(key,"")
            if val:
                label=item.get("heading") or default_label
                if item.get("show_heading", True): doc.add_heading(label, level=1)
                _add_body_paragraph(doc,val); rendered_roles.add(role)
    # Fallback for journals that do not prescribe declaration order.
    for role,(label,key) in role_map.items():
        if role in rendered_roles: continue
        val=pkg.get(key,"")
        if val:
            doc.add_heading(label,level=1); _add_body_paragraph(doc,val)

    refs = pkg.get("references") or []
    if refs:
        ref_heading="References"; ref_show=True
        for item in outline:
            if item.get("role")=="references":
                ref_heading=item.get("heading") or ref_heading; ref_show=item.get("show_heading",True); break
        if ref_show: doc.add_heading(ref_heading, level=1)
        for i, ref in enumerate(refs, 1):
            ref = re.sub(r"^\s*\d+[\.)]?\s+", "", str(ref))
            p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Inches(-0.2); p.paragraph_format.left_indent = Inches(0.2)
            p.add_run(f"{i}. ").bold = True; p.add_run(ref)

    # Extended Data items included in the same initial-submission manuscript when requested.
    ext_figs = [f for f in pkg.get("figure_plan",[]) if f.get("destination") == "extended_data" and f.get("action") != "remove_if_redundant"]
    ext_tables = [t for t in pkg.get("tables",[]) if t.get("destination") == "extended_data"]
    if ext_figs or ext_tables:
        doc.add_page_break(); doc.add_heading("Extended Data", level=1)
        for t in ext_tables: _insert_table(doc, t)
        for f in ext_figs: _insert_figure(doc, f, source_images, redrawn_images)

    bio = BytesIO(); doc.save(bio); return bio.getvalue()


def cover_letter_docx_bytes(pkg: dict, journal: str) -> bytes:
    import copy
    pkg = _clean_pkg_in_place(copy.deepcopy(pkg))
    doc = Document(); _style_doc(doc)
    cl = pkg.get("cover_letter") or {}
    doc.add_heading("Cover Letter", level=0)
    doc.add_paragraph(f"To: Editors, {journal}")
    if cl.get("subject"):
        p=doc.add_paragraph(); p.add_run("Re: ").bold=True; p.add_run(cl["subject"])
    doc.add_paragraph(cl.get("salutation") or "Dear Editors,")
    for para in cl.get("paragraphs",[]) or []: _add_body_paragraph(doc, para)
    doc.add_paragraph(cl.get("closing") or "Sincerely,")
    doc.add_paragraph(cl.get("signatory") or "Corresponding authors")
    bio=BytesIO(); doc.save(bio); return bio.getvalue()


def supplementary_docx_bytes(pkg: dict, *, source_images: list[dict], redrawn_images: dict[int, bytes] | None = None) -> bytes | None:
    import copy
    pkg = _clean_pkg_in_place(copy.deepcopy(pkg))
    secs = pkg.get("supplementary_sections",[]) or []
    tabs = [t for t in pkg.get("tables",[]) if t.get("destination") == "supplementary"]
    figs = [f for f in pkg.get("figure_plan",[]) if f.get("destination") == "supplementary" and f.get("action") != "remove_if_redundant"]
    if not (secs or tabs or figs): return None
    doc=Document(); _style_doc(doc); doc.add_heading("Supplementary Information", level=0); doc.add_paragraph(pkg.get("title", ""))
    for sec in secs:
        if sec.get("show_heading") and sec.get("heading"): doc.add_heading(sec["heading"],level=1)
        for para in sec.get("paragraphs",[]) or []: _add_body_paragraph(doc,para)
    for t in tabs: _insert_table(doc,t)
    for f in figs: _insert_figure(doc,f,source_images,redrawn_images)
    bio=BytesIO(); doc.save(bio); return bio.getvalue()


def report_docx_bytes(journal_profile: dict, figure_audit: dict, reference_audit: dict, review: dict, pkg: dict) -> bytes:
    doc = Document(); _style_doc(doc)
    doc.add_heading("NanoMax Nature Editor — Submission Audit", level=0)
    doc.add_paragraph(f"Target: {journal_profile.get('journal','')} — {journal_profile.get('article_type','')}")
    doc.add_heading("Final submission gate", level=1)
    doc.add_paragraph(f"Decision: {review.get('editorial_decision','')} | Readiness: {review.get('submission_readiness_score',0)}/100")
    for label,key in [("Blocking issues","blocking_issues"),("Major issues","major_issues"),("Minor issues","minor_issues"),("Strengths","strengths"),("Final actions","final_actions")]:
        vals=review.get(key,[]) or []
        if vals:
            doc.add_heading(label,level=2)
            for x in vals: doc.add_paragraph(x,style="List Bullet")
    doc.add_heading("Author actions", level=1)
    for x in pkg.get("author_actions",[]) or []:
        doc.add_paragraph(f"[{x.get('severity','')}] {x.get('location','')}: {x.get('issue','')} — {x.get('required_action','')}", style="List Bullet")
    doc.add_heading("Figure integrity audit", level=1); doc.add_paragraph(figure_audit.get("summary",""))
    for a in figure_audit.get("assets",[]) or []:
        doc.add_paragraph(f"Asset {a.get('asset_index')}: {a.get('recommended_action')} — {a.get('reason')}",style="List Bullet")
    doc.add_heading("Reference/citation audit", level=1); doc.add_paragraph(reference_audit.get("summary",""))
    for x in reference_audit.get("needs_author_check",[]) or []: doc.add_paragraph(x,style="List Bullet")
    bio=BytesIO(); doc.save(bio); return bio.getvalue()


def package_zip_bytes(*, master_docx: bytes, cover_letter_docx: bytes, report_docx: bytes, transformed: dict, journal_profile: dict, figure_audit: dict, reference_audit: dict, review: dict, source_images: list[dict], redrawn_images: dict[int, bytes] | None = None, graphical_abstract: bytes | None = None, supplementary_docx: bytes | None = None) -> bytes:
    out=BytesIO()
    with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
        z.writestr("01_Submission/NanoMax_Nature_Editor_Final_Manuscript.docx", master_docx)
        z.writestr("01_Submission/Cover_Letter.docx", cover_letter_docx)
        if supplementary_docx: z.writestr("01_Submission/Supplementary_Information.docx", supplementary_docx)
        if graphical_abstract: z.writestr("01_Submission/Graphical_Abstract.png", graphical_abstract)
        z.writestr("02_Quality_Control/NanoMax_Submission_Audit.docx", report_docx)
        z.writestr("02_Quality_Control/session.json", json.dumps({"journal_profile":journal_profile,"figure_audit":figure_audit,"reference_audit":reference_audit,"review":review,"final_manuscript":transformed},ensure_ascii=False,indent=2))
        for i,img in enumerate(source_images,1): z.writestr(f"03_Source_Figures/FigureAsset_{i}.{img.get('ext','png')}",img["bytes"])
        for idx,blob in (redrawn_images or {}).items(): z.writestr(f"04_Redrawn_Figures/FigureAsset_{idx}_redrawn.png",blob)
    return out.getvalue()


def research_completion_docx_bytes(plan: dict) -> bytes:
    doc=Document(); _style_doc(doc)
    doc.add_heading("NanoMax Research Completion Plan", level=0)
    doc.add_paragraph(plan.get("summary", ""))
    p=doc.add_paragraph(); r=p.add_run(plan.get("synthetic_boundary_statement", "SIMULATED / HYPOTHETICAL outputs are planning artifacts only.")); r.bold=True
    for gap in plan.get("gaps",[]) or []:
        doc.add_heading(f"{gap.get('gap_id','Gap')} — {gap.get('priority','').title()}", level=1)
        for label,key in [
            ("Manuscript location","manuscript_location"),("Missing evidence","missing_evidence"),("Why it matters","why_needed"),
            ("Proposed experiment or analysis","proposed_experiment_or_analysis"),("Claim potentially supported if real evidence succeeds","claim_unlocked_if_real")]:
            val=gap.get(key,"")
            if val:
                p=doc.add_paragraph(); p.add_run(label+": ").bold=True; p.add_run(str(val))
        for label,key in [("Required inputs","required_inputs"),("Protocol","protocol_steps"),("Controls","controls"),("Statistical plan","statistical_plan"),("Expected real outputs","real_output_expected")]:
            vals=gap.get(key,[]) or []
            if vals:
                doc.add_heading(label, level=2)
                for x in vals: doc.add_paragraph(str(x), style="List Bullet")
        if gap.get("synthetic_preview_type") != "none":
            doc.add_heading("Synthetic preview (planning only)", level=2)
            doc.add_paragraph(f"Type: {gap.get('synthetic_preview_type')}")
            doc.add_paragraph(gap.get("synthetic_preview_prompt", ""))
    bio=BytesIO(); doc.save(bio); return bio.getvalue()


def submission_zip_bytes(*, master_docx: bytes, cover_letter_docx: bytes, supplementary_docx: bytes | None = None, graphical_abstract: bytes | None = None) -> bytes:
    """Journal-facing files only. No NanoMax audit/session/synthetic material."""
    out=BytesIO()
    with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
        z.writestr("Final_Manuscript.docx", master_docx)
        z.writestr("Cover_Letter.docx", cover_letter_docx)
        if supplementary_docx: z.writestr("Supplementary_Information.docx", supplementary_docx)
        if graphical_abstract: z.writestr("Graphical_Abstract.png", graphical_abstract)
    return out.getvalue()


def internal_qc_zip_bytes(*, report_docx: bytes, transformed: dict, journal_profile: dict, figure_audit: dict, reference_audit: dict, review: dict, source_images: list[dict], redrawn_images: dict[int, bytes] | None = None, research_plan_docx: bytes | None = None, research_plan: dict | None = None, synthetic_previews: dict[str, bytes] | None = None) -> bytes:
    out=BytesIO()
    with zipfile.ZipFile(out,"w",zipfile.ZIP_DEFLATED) as z:
        z.writestr("NanoMax_Submission_Audit.docx", report_docx)
        z.writestr("session.json", json.dumps({"journal_profile":journal_profile,"figure_audit":figure_audit,"reference_audit":reference_audit,"review":review,"final_manuscript":transformed,"research_completion":research_plan or {}},ensure_ascii=False,indent=2))
        if research_plan_docx: z.writestr("Research_Completion_Plan.docx", research_plan_docx)
        for i,img in enumerate(source_images,1): z.writestr(f"Source_Figures/FigureAsset_{i}.{img.get('ext','png')}",img["bytes"])
        for idx,blob in (redrawn_images or {}).items(): z.writestr(f"Redrawn_Conceptual_Figures/FigureAsset_{idx}_redrawn.png",blob)
        for key,blob in (synthetic_previews or {}).items(): z.writestr(f"Synthetic_Previews/{key}.png",blob)
    return out.getvalue()
