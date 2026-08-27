from io import BytesIO
from docx import Document
from PIL import Image
from nanomax_editor.exporters import master_manuscript_docx_bytes, cover_letter_docx_bytes


def _png():
    bio=BytesIO(); Image.new("RGB",(800,450),"white").save(bio,"PNG"); return bio.getvalue()


def _pkg():
    return {
        "title":"A concise scientific title","short_title":"Short",
        "front_matter":{"authors_line":"A. Author1, B. Author2*","affiliations":["1 Institute A","2 Institute B"],"correspondence":"* email@example.org","equal_contribution":"","other_notes":[]},
        "abstract":"A complete abstract with [[CITE:1]] only for exporter testing.","keywords":[],
        "main_sections":[{"heading":"Introduction","show_heading":False,"paragraphs":["Intro text [[CITE:1,2]]."]},{"heading":"Results","show_heading":True,"paragraphs":["Result text."]},{"heading":"Discussion","show_heading":True,"paragraphs":["Discussion text."]}],
        "methods_sections":[{"heading":"Study design","show_heading":True,"paragraphs":["Methods text."]}],
        "tables":[{"number":"Table 1","title":"Test table","columns":["A","B"],"rows":[["1","2"]],"footnote":"Footnote.","placement_after":"Results","destination":"main"}],
        "figure_plan":[{"figure_number":"Fig. 1","source_asset_indices":[1],"action":"preserve","title":"Test figure","legend":"Legend text.","placement_after":"Results","destination":"main","reason":"test","redraw_prompt":""}],
        "references":["Author A. Paper title. Journal 1, 1–2 (2026).","Author B. Paper two. Journal 2, 3–4 (2025)."],
        "data_availability":"Data statement.","code_availability":"Code statement.","ethics_statement":"Ethics statement.","author_contributions":"Contributions.","competing_interests":"No competing interests.","funding_statement":"Funding.","acknowledgements":"Thanks.",
        "graphical_abstract":{"recommended":False,"include_in_manuscript":False,"source_asset_index":0,"reason":"","generation_prompt":"","caption":""},
        "cover_letter":{"subject":"Submission of A concise scientific title","salutation":"Dear Editors,","paragraphs":["Please consider our manuscript."],"closing":"Sincerely,","signatory":"A. Author"},
        "supplementary_sections":[],"author_actions":[],"change_summary":[]
    }


def test_master_contains_table_and_image():
    data=master_manuscript_docx_bytes(_pkg(),source_images=[{"bytes":_png(),"name":"f.png","ext":"png"}])
    doc=Document(BytesIO(data))
    assert len(doc.tables) == 1
    assert len(doc.inline_shapes) == 1
    text="\n".join(p.text for p in doc.paragraphs)
    assert "A concise scientific title" in text
    assert "AUTHOR_INPUT_NEEDED" not in text


def test_cover_letter():
    data=cover_letter_docx_bytes(_pkg(),"Nature Cancer")
    doc=Document(BytesIO(data))
    assert "Nature Cancer" in "\n".join(p.text for p in doc.paragraphs)
